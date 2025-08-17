"""
Unit tests for file processing service.
"""

import pytest
import tempfile
import os
from unittest.mock import Mock, patch, AsyncMock

from fastapi import UploadFile, HTTPException

from src.services.file_processing import FileProcessingService


class TestFileProcessingService:
    """Test cases for FileProcessingService."""

    @pytest.mark.asyncio
    async def test_process_txt_file_success(self):
        """Test successful .txt file processing."""
        
        content = "This is a test transcript that is definitely longer than one hundred characters to pass validation."
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write(content)
            temp_file_path = f.name
        
        try:
            # Create mock UploadFile
            with open(temp_file_path, 'rb') as file:
                mock_upload_file = Mock(spec=UploadFile)
                mock_upload_file.filename = "test.txt"
                mock_upload_file.read = AsyncMock(return_value=file.read())
                mock_upload_file.seek = AsyncMock()
                
                result = await FileProcessingService.process_uploaded_file(mock_upload_file)
                
                assert result.strip() == content
                
        finally:
            os.unlink(temp_file_path)

    @pytest.mark.asyncio
    async def test_process_docx_file_success(self):
        """Test successful .docx file processing."""
        
        # Create a simple docx file for testing
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("This is a test paragraph in a DOCX file.")
        doc.add_paragraph("This content should be long enough to pass validation checks.")
        doc.add_paragraph("The file processing service should extract this text correctly.")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_file_path = f.name
        
        try:
            with open(temp_file_path, 'rb') as file:
                mock_upload_file = Mock(spec=UploadFile)
                mock_upload_file.filename = "test.docx"
                mock_upload_file.read = AsyncMock(return_value=file.read())
                mock_upload_file.seek = AsyncMock()
                
                result = await FileProcessingService.process_uploaded_file(mock_upload_file)
                
                assert "test paragraph in a DOCX file" in result
                assert "content should be long enough" in result
                assert len(result) > 100  # Should pass length validation
                
        finally:
            os.unlink(temp_file_path)

    @pytest.mark.asyncio
    async def test_file_too_large(self):
        """Test file size validation."""
        
        # Create a large content string
        large_content = "A" * (FileProcessingService.MAX_FILE_SIZE + 1)
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "large.txt"
        mock_upload_file.read = AsyncMock(return_value=large_content.encode())
        mock_upload_file.seek = AsyncMock()
        
        with pytest.raises(HTTPException) as exc_info:
            await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert exc_info.value.status_code == 413
        assert "too large" in exc_info.value.detail.lower()

    @pytest.mark.asyncio
    async def test_unsupported_file_extension(self):
        """Test unsupported file type validation."""
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "test.pdf"
        mock_upload_file.read = AsyncMock(return_value=b"test content")
        mock_upload_file.seek = AsyncMock()
        
        with pytest.raises(HTTPException) as exc_info:
            await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert exc_info.value.status_code == 400
        assert "unsupported" in exc_info.value.detail.lower()

    @pytest.mark.asyncio
    async def test_no_filename(self):
        """Test error when no filename is provided."""
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = None
        mock_upload_file.read = AsyncMock(return_value=b"test content")
        mock_upload_file.seek = AsyncMock()
        
        with pytest.raises(HTTPException) as exc_info:
            await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert exc_info.value.status_code == 400
        assert "filename" in exc_info.value.detail.lower()

    @pytest.mark.asyncio
    async def test_content_too_short(self):
        """Test content length validation."""
        
        short_content = "Short"  # Less than 100 characters
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "short.txt"
        mock_upload_file.read = AsyncMock(return_value=short_content.encode())
        mock_upload_file.seek = AsyncMock()
        
        with pytest.raises(HTTPException) as exc_info:
            await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert exc_info.value.status_code == 400
        assert "too short" in exc_info.value.detail.lower()

    @pytest.mark.asyncio
    async def test_content_too_large(self):
        """Test content size validation."""
        
        # Create content larger than MAX_TEXT_SIZE
        large_content = "A" * (FileProcessingService.MAX_TEXT_SIZE + 1)
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "large.txt"
        mock_upload_file.read = AsyncMock(return_value=large_content.encode())
        mock_upload_file.seek = AsyncMock()
        
        with pytest.raises(HTTPException) as exc_info:
            await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert exc_info.value.status_code == 413
        assert "content too large" in exc_info.value.detail.lower()

    @pytest.mark.asyncio
    async def test_malicious_content_detection(self):
        """Test malicious content detection."""
        
        malicious_content = "This transcript contains <script>alert('xss')</script> and some normal coaching content."
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "malicious.txt"
        mock_upload_file.read = AsyncMock(return_value=malicious_content.encode())
        mock_upload_file.seek = AsyncMock()
        
        with pytest.raises(HTTPException) as exc_info:
            await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert exc_info.value.status_code == 400
        assert "malicious" in exc_info.value.detail.lower()

    @pytest.mark.asyncio
    async def test_encoding_detection_utf8(self):
        """Test UTF-8 encoding detection."""
        
        content = "Test content with UTF-8 characters: éñíñé"
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "utf8.txt"
        mock_upload_file.read = AsyncMock(return_value=content.encode('utf-8'))
        mock_upload_file.seek = AsyncMock()
        
        result = await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert "éñíñé" in result

    @pytest.mark.asyncio
    async def test_encoding_detection_latin1(self):
        """Test Latin-1 encoding detection and handling."""
        
        content = "Test content with Latin-1 characters"
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "latin1.txt"
        mock_upload_file.read = AsyncMock(return_value=content.encode('latin-1'))
        mock_upload_file.seek = AsyncMock()
        
        # Mock chardet to return latin-1 encoding
        with patch('src.services.file_processing.chardet.detect') as mock_detect:
            mock_detect.return_value = {'encoding': 'latin-1'}
            
            result = await FileProcessingService.process_uploaded_file(mock_upload_file)
            
            assert content in result

    @pytest.mark.asyncio
    async def test_docx_processing_error(self):
        """Test error handling for corrupted DOCX files."""
        
        # Create invalid DOCX content
        invalid_docx_content = b"This is not a valid DOCX file"
        
        mock_upload_file = Mock(spec=UploadFile)
        mock_upload_file.filename = "invalid.docx"
        mock_upload_file.read = AsyncMock(return_value=invalid_docx_content)
        mock_upload_file.seek = AsyncMock()
        
        with pytest.raises(HTTPException) as exc_info:
            await FileProcessingService.process_uploaded_file(mock_upload_file)
        
        assert exc_info.value.status_code == 422
        assert "failed to process" in exc_info.value.detail.lower()

    def test_get_file_extension(self):
        """Test file extension extraction."""
        
        assert FileProcessingService._get_file_extension("test.txt") == ".txt"
        assert FileProcessingService._get_file_extension("test.DOCX") == ".docx"
        assert FileProcessingService._get_file_extension("file.with.dots.txt") == ".txt"
        
        with pytest.raises(HTTPException):
            FileProcessingService._get_file_extension("")
        
        with pytest.raises(HTTPException):
            FileProcessingService._get_file_extension(None)

    def test_validate_file_extension(self):
        """Test file extension validation."""
        
        # Valid extensions should not raise
        FileProcessingService._validate_file_extension(".txt")
        FileProcessingService._validate_file_extension(".docx")
        
        # Invalid extensions should raise
        with pytest.raises(HTTPException) as exc_info:
            FileProcessingService._validate_file_extension(".pdf")
        
        assert exc_info.value.status_code == 400
        assert "unsupported" in exc_info.value.detail.lower()

    def test_scan_for_malicious_content(self):
        """Test malicious content scanning."""
        
        # Safe content should not raise
        safe_content = "This is a normal coaching session transcript."
        FileProcessingService._scan_for_malicious_content(safe_content)
        
        # Malicious patterns should raise
        malicious_patterns = [
            "This contains <script>alert('xss')</script>",
            "javascript:void(0)",
            "eval(malicious_code)",
            "document.cookie = 'stolen'",
            "window.location = 'evil.com'"
        ]
        
        for pattern in malicious_patterns:
            with pytest.raises(HTTPException) as exc_info:
                FileProcessingService._scan_for_malicious_content(pattern)
            
            assert exc_info.value.status_code == 400
            assert "malicious" in exc_info.value.detail.lower()